"""
任务报告生成 — 把 task + artifacts 渲染成 .docx
聚焦数据分析与形变结论,不做过程性描述。

入口: build_task_report(task, artifacts, output_path)
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)


def _kv_table(doc, rows: list[tuple[str, Any]]):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells
        c0.text = str(k)
        c1.text = "-" if v is None or v == "" else str(v)
        for c in (c0, c1):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in c0.paragraphs[0].runs:
            run.bold = True


def _caption(doc, text: str):
    p = doc.add_paragraph(text)
    p.alignment = 1
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)


def _classify_aoi(aoi_name: str) -> str:
    """从 AOI 名称识别业务场景,影响后续业务解读用语。"""
    n = (aoi_name or "").lower()
    if any(k in aoi_name for k in ("矿区", "金矿", "煤矿", "铜矿", "铁矿", "采空区", "矿山")):
        return "mining"
    if any(k in aoi_name for k in ("油田", "气田", "油气", "抽水")):
        return "oilgas"
    if any(k in aoi_name for k in ("断裂", "断层", "构造")):
        return "tectonic"
    if any(k in aoi_name for k in ("滑坡", "边坡", "崩塌")):
        return "landslide"
    if any(k in aoi_name for k in ("地热", "火山")):
        return "geothermal"
    return "general"


def _build_conclusion(aoi_name: str, v_stats: Dict | None,
                      coh_overall: float | None,
                      time_span_days: int | None) -> Dict[str, Any]:
    """
    生成针对性结论。返回 dict:
      verdict        : 一句话核心判断
      highlights     : list[str] 关键量化观察
      interpretation : 业务/地质解读(基于 AOI 类型)
      recommendations: list[str] 具体行动建议
      confidence     : 置信度等级
    """
    scene = _classify_aoi(aoi_name)
    scene_label = {
        "mining": "矿区沉降监测", "oilgas": "油气田抽采监测",
        "tectonic": "构造活动监测", "landslide": "边坡稳定性监测",
        "geothermal": "地热/火山活动监测", "general": "一般地表形变监测",
    }[scene]

    if not v_stats or v_stats.get("p5") is None:
        return {
            "verdict": "未生成 SBAS 反演结果,无法给出形变判断。",
            "highlights": [],
            "interpretation": "",
            "recommendations": ["在详情页点击「跑 SBAS」生成反演结果后重新出报告。"],
            "confidence": "—",
            "scene_label": scene_label,
        }

    p5  = v_stats["p5"]
    p95 = v_stats["p95"]
    std = v_stats.get("std") or 0
    vmean = v_stats.get("mean") or 0
    vmin  = v_stats.get("min", 0)
    vmax  = v_stats.get("max", 0)
    span = p95 - p5
    span_yr = (time_span_days / 365.25) if time_span_days else None

    # 判定主信号强度(基于稳健的 P5/P95,不被极值 outlier 干扰)
    significant_subsidence = p5 < -10
    significant_uplift     = p95 > 15
    stable = abs(p5) < 8 and abs(p95) < 10 and std < 20

    highlights = []
    highlights.append(
        f"LOS 形变速率 95% 区间(P5~P95): **[{p5:+.1f}, {p95:+.1f}] mm/yr**,跨度 {span:.1f} mm/yr"
    )
    highlights.append(
        f"速率均值 {vmean:+.2f} mm/yr,标准差 {std:.2f} mm/yr"
    )
    if abs(vmin) > 100 or abs(vmax) > 100:
        highlights.append(
            f"极值范围 {vmin:+.1f} ~ {vmax:+.1f} mm/yr,显著超出 P5~P95 区间 — "
            f"判定为单像素 unwrap 错误或大气湍流 artifacts,不代表真实地质形变"
        )
    if coh_overall is not None:
        coh_level = "优秀(≥0.7)" if coh_overall >= 0.7 else ("良好(0.4~0.7)" if coh_overall >= 0.4 else "偏低(<0.4)")
        highlights.append(
            f"整体平均相干性 {coh_overall:.2f},数据质量 {coh_level}"
        )

    # 核心判断 + 业务解读
    if stable:
        verdict = (
            f"**本研究区({aoi_name})整体地表稳定,未观测到 ≥ 10 mm/yr 量级的显著形变信号。**"
            f"95% 像素速率分布集中在 [{p5:+.1f}, {p95:+.1f}] mm/yr,处于本反演的噪声水平内,"
            f"可判定为无显著地质形变活动。"
        )
        confidence = "中-高" if coh_overall and coh_overall >= 0.7 else "中"
        interp = {
            "mining": (
                "矿区形变背景稳定,无明显采动活动信号,符合成熟矿区/闭坑后/未开采状态特征。\n\n"
                "**对矿权决策的辅助价值**:\n"
                "  • 若本区为勘探候选/未开采:可作为后续探矿权申请、钻探施工的「形变基线」存档,"
                "便于将来开采阶段对比识别新增沉降\n"
                "  • 若本区已知有矿权:稳定结果可佐证开采节奏低或采空区不在本 AOI 范围内;"
                "若本应有活跃开采却无沉降信号,需进一步核查实际生产状态\n\n"
                "**重要免责**:InSAR 仅观测地表形变,与地下矿种类型、储量品位**直接无关**。"
                "矿种识别需结合高光谱(蚀变矿物)/光学(岩性)/地球物理(深部构造)等独立证据,"
                "本报告不能用于判断「该区是否具备某矿种储量」。"
            ),
            "oilgas": "油气田未观测到典型抽采沉降漏斗,可能反映**注水回灌平衡或开发强度低**。",
            "tectonic": "构造带未观测到显著蠕动信号,符合**间震期/已锁定段**特征,与年际背景一致。",
            "landslide": "边坡未观测到加速变形,处于**稳定阶段**。",
            "geothermal": "未观测到显著膨胀/收缩信号,**热储压力变化平缓**。",
            "general": "AOI 范围地表稳定,无需特别关注。",
        }[scene]
        recommendations = [
            f"持续监测:每 6 个月重跑一次时序反演,对比是否出现新增形变信号。",
            f"如对**毫米级**精度有要求(典型场景:{scene_label}的早期预警),建议接入 GACOS 大气校正,"
            f"可把绝对速率不确定度从 ±5 mm/yr 压到 ±1~2 mm/yr。",
        ]
    elif significant_subsidence:
        verdict = (
            f"**本研究区({aoi_name})存在沉降迹象**:5% 像素速率低于 {p5:.1f} mm/yr,"
            f"P5~P95 跨度 {span:.1f} mm/yr 超出典型噪声水平({'数据相干性高,可信度较好' if coh_overall and coh_overall>=0.7 else '需结合相干性进一步判读'})。"
        )
        confidence = "中"
        interp = {
            "mining": (
                "检测到沉降迹象,符合采动型形变(掘进掌子面/塌陷漏斗)的典型空间特征。\n\n"
                "**对矿权决策的辅助价值**:\n"
                "  • 若本区已知为在产/历史矿区:沉降漏斗位置可用于**验证已知采空区边界**,"
                "或**定位未报备的盗采点**(漏斗与已申报矿权范围不重合时尤其有价值)\n"
                "  • 若本区原本应为空白区(无报备矿权):出现采动型沉降可能指示**非法采矿活动**,建议地面核查\n"
                "  • 真伪判别要点:沉降的连片性 + 时序单调性。连片漏斗 + 单调累积 = 真采动信号;"
                "孤立像素或阶梯跳变 = 大气/unwrap artifacts\n\n"
                "**重要免责**:沉降信号能证明「有采矿活动或自然塌陷」,但**不能反推矿种类型或储量**。"
                "矿种识别仍需高光谱/光学/地球物理独立证据。"
            ),
            "oilgas": "符合**抽采沉降**典型特征,沉降中心通常对应主力生产井位置。",
            "tectonic": "可能对应**断层蠕滑**或**地下水位下降导致的区域性沉降**,需结合区域地质背景判读。",
            "landslide": "**警惕**:加速变形是边坡失稳前兆,需结合降雨/施工等触发因素综合判断,必要时启动现场监测。",
            "geothermal": "可能反映**热储压力下降**或**注水井附近的局部沉降**。",
            "general": "存在局部沉降信号,需结合区域地质背景判读其性质。",
        }[scene]
        recommendations = [
            f"**先看速率图空间分布**:沉降是连片区域(geologically meaningful)还是孤立像素(可能是 artifacts)?连片 → 真实信号;孤立 → 可能 unwrap 错误。",
            f"**做时序剖面**:在沉降疑似区域选 3-5 个点,看累积形变曲线是否单调下降。若是阶梯状或大幅震荡,大概率是 artifacts。",
            f"**接 GACOS 大气校正后重做**,确认信号在去除大气残差后仍然存在。",
        ]
    elif significant_uplift:
        verdict = (
            f"**本研究区({aoi_name})存在异常正向形变信号**:5% 像素速率高于 {p95:.1f} mm/yr。"
            f"此量级在典型 {scene_label} 场景中较罕见,**优先怀疑为残余大气延迟或 unwrap 错误污染**。"
        )
        confidence = "低"
        interp = "正向形变(地表抬升)在地质监测里较罕见的真实场景包括火山膨胀、深层注水、冻胀融沉。其余多为 InSAR artifacts。"
        recommendations = [
            f"在抬升疑似区选点做时序剖面,若曲线含大跳变(≥ 28mm 的整数倍)即可确认是 unwrap 错误。",
            f"接 GACOS 校正后重做反演,残余大气是主要嫌疑。",
        ]
    else:
        verdict = (
            f"**本研究区({aoi_name})存在缓慢形变趋势**:P5~P95 = [{p5:+.1f}, {p95:+.1f}] mm/yr,"
            f"信号在大气噪声地板边缘,本次反演不能确证。"
        )
        confidence = "中(信号在噪声边缘)"
        if scene == "mining":
            interp = (
                f"速率分布跨度 {span:.1f} mm/yr,信号在大气噪声边缘。"
                f"既可能是缓慢采动沉降(深部开采/小规模盗采),也可能是未校正的大气延迟残差。\n\n"
                f"**对矿权决策的辅助价值**:本次结果**不足以**支撑「该区是否有采矿活动」的可靠判断,"
                f"建议接 GACOS 大气校正后重做。若仍存在弱信号,可作为「疑似采动迹象」的初筛输入。\n\n"
                f"**重要免责**:即便确认有采矿活动,InSAR 也不能告诉你矿种或储量。"
            )
        else:
            interp = (
                f"速率分布跨度 {span:.1f} mm/yr,既可能是真实的缓慢区域沉降(典型 {scene_label} 慢速过程),"
                f"也可能是未校正的大气延迟残差。两者在本反演中无法区分。"
            )
        recommendations = [
            f"**接 GACOS 大气校正后重做反演** —— 这是当前的关键约束,可把噪声地板从 ±15 mm/yr 压到 ±3 mm/yr,真实信号会浮出来。",
            f"扩长观测时段(≥ 2 年),延长基线让弱信号在累积形变上更显著。",
        ]

    return {
        "verdict": verdict,
        "highlights": highlights,
        "interpretation": interp,
        "recommendations": recommendations,
        "confidence": confidence,
        "scene_label": scene_label,
    }


def _set_chinese_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def build_task_report(task: Dict, artifacts: Dict, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_chinese_font(doc)

    # ── 封面 ────────────────────────────────────
    doc.add_heading(f"InSAR 形变监测分析报告", level=0)
    sub = doc.add_paragraph()
    sub.add_run(task.get("label") or task.get("aoi_name") or "(无标题)").bold = True
    sub.add_run(f"   ·   报告生成 {datetime.utcnow().strftime('%Y-%m-%d')}").italic = True

    # ── 1. AOI 与观测概况 ──────────────────────
    _heading(doc, "1. AOI 与观测概况")
    progress = task.get("progress") or {}
    stack = artifacts.get("stack_index") or {}
    dr = stack.get("date_range") or [task.get("start_date"), task.get("end_date")]
    _kv_table(doc, [
        ("研究区",       task.get("aoi_name")),
        ("观测时段",     f"{dr[0]}  →  {dr[1]}"),
        ("数据源",       "Sentinel-1 C-band SAR(欧空局)"),
        ("有效干涉对数", stack.get("pair_count") or f"{progress.get('done', 0)} 对"),
        ("极化",         ", ".join(stack.get("polarizations") or [task.get("polarization", "VV")])),
        ("轨道方向",     ", ".join(stack.get("orbit_directions") or [])),
    ])

    # ── 2. 数据特征 ────────────────────────────
    _heading(doc, "2. 数据特征")

    coh_overall = stack.get("coherence_mean_overall")
    if stack and not stack.get("_error"):
        bsl_list = sorted(set(stack.get("temporal_baselines_days") or []))
        rows = [
            ("整体平均相干性", f"{coh_overall:.3f}" if coh_overall is not None else "-"),
            ("时间基线分布", f"{bsl_list[0]}~{bsl_list[-1]} 天" if bsl_list else "-"),
            ("典型基线值",   f"{', '.join(map(str, bsl_list[:5]))} 天" if bsl_list else "-"),
        ]
        _kv_table(doc, rows)
        if coh_overall is not None:
            p = doc.add_paragraph()
            p.add_run("相干性解读:").bold = True
            if coh_overall >= 0.7:
                p.add_run(f" 平均相干性 {coh_overall:.2f} 处于高水平,"
                          f"表明 AOI 地表稳定(植被/积水/扰动少),InSAR 信号质量良好。")
            elif coh_overall >= 0.4:
                p.add_run(f" 平均相干性 {coh_overall:.2f} 处于中等水平,部分区域可能受植被或地表扰动影响。")
            else:
                p.add_run(f" 平均相干性 {coh_overall:.2f} 偏低,反演结果不确定性较大。")

    # ── 2.1 代表性快视图 ───────────────────────
    quicklooks = artifacts.get("quicklooks") or []
    if quicklooks:
        _heading(doc, "2.1 代表性干涉对快视", level=2)
        doc.add_paragraph(
            f"以下展示 {min(len(quicklooks), 6)} 张相干性最高的干涉对快视图,"
            f"每张包含 LOS 形变(中值去趋势)与相干性两个子图。"
            f"形变图色标:红色 = 远离卫星方向(可能对应沉降),蓝色 = 靠近卫星方向(可能对应抬升)。"
        )
        for ql in quicklooks[:6]:
            png = _resolve_files_url(ql.get("url"))
            if png and png.exists():
                try:
                    doc.add_picture(str(png), width=Cm(16))
                    _caption(doc, ql.get("name", ""))
                except Exception as e:
                    doc.add_paragraph(f"  [插图失败] {ql.get('name')}: {e}")

    # ── 3. 形变速率分析 ───────────────────────
    _heading(doc, "3. 形变速率分析")
    sbas_list = artifacts.get("sbas") or []
    if not sbas_list:
        doc.add_paragraph(
            "未生成 SBAS 时序反演结果。无法给出形变速率与时序判断。",
            style="Intense Quote",
        )
    else:
        doc.add_paragraph(
            "采用简化 SBAS(Small Baseline Subset)时序反演,"
            "将所有干涉对的 LOS 相位联合反演为每个 SAR 时相的累积形变,"
            "并由其时间序列拟合斜率得到 LOS 速率(mm/yr)。"
        )

    for idx, sb in enumerate(sbas_list, 1):
        _heading(doc, f"3.{idx} burst {sb['burst']}", level=2)
        sm = sb.get("summary") or {}
        v = sm.get("velocity_mm_per_year") or {}

        _kv_table(doc, [
            ("观测时段",        " → ".join(sm.get("date_range") or [])),
            ("使用 pair 数",     f"{sm.get('n_pairs', '-')} 对  /  {sm.get('n_dates', '-')} 时相"),
            ("有效像素覆盖",     f"{sm.get('valid_pixel_pct', 0):.1f}%" if sm.get("valid_pixel_pct") is not None else "-"),
            ("速率统计 (mm/yr)", f"min {v.get('min', 0):.1f}   max {v.get('max', 0):.1f}   均值 {v.get('mean', 0):.2f}" if v else "-"),
            ("**P5 ~ P95 区间**", f"**[{v.get('p5', 0):.1f}, {v.get('p95', 0):.1f}] mm/yr**" if v else "-"),
            ("标准差",          f"{v.get('std', 0):.2f} mm/yr" if v else "-"),
        ])

        vel_path = _resolve_files_url(sb.get("velocity_map_url"))
        if vel_path and vel_path.exists():
            doc.add_picture(str(vel_path), width=Cm(15))
            _caption(doc, "LOS 速率图(空间分布)")
        ts_path = _resolve_files_url(sb.get("timeseries_points_url"))
        if ts_path and ts_path.exists():
            doc.add_picture(str(ts_path), width=Cm(16))
            _caption(doc, "代表点位累积形变时序曲线(P1 最大正速率 / P2 最大负速率 / P3 稳定点)")

    # ── 4. 结论与建议 ────────────────────────
    _heading(doc, "4. 结论与建议")
    first_sbas = sbas_list[0] if sbas_list else None
    v_stats = ((first_sbas.get("summary") or {}).get("velocity_mm_per_year")
               if first_sbas else None)

    # 计算观测时长
    time_span_days = None
    if stack and stack.get("date_range"):
        try:
            d0 = datetime.fromisoformat(stack["date_range"][0])
            d1 = datetime.fromisoformat(stack["date_range"][1])
            time_span_days = (d1 - d0).days
        except Exception:
            pass

    concl = _build_conclusion(task.get("aoi_name", ""), v_stats, coh_overall, time_span_days)

    # 4.1 核心判断(一句话,加粗)
    _heading(doc, "4.1 核心判断", level=2)
    p = doc.add_paragraph()
    p.add_run(concl["verdict"])

    # 4.2 关键观察(量化 bullets)
    _heading(doc, "4.2 关键观察", level=2)
    for h in concl["highlights"]:
        doc.add_paragraph(h, style="List Bullet")

    # 4.3 业务解读(基于 AOI 类型)
    if concl["interpretation"]:
        _heading(doc, f"4.3 业务解读({concl['scene_label']})", level=2)
        doc.add_paragraph(concl["interpretation"])

    # 4.4 下一步建议
    _heading(doc, "4.4 下一步建议", level=2)
    for r in concl["recommendations"]:
        doc.add_paragraph(r, style="List Bullet")

    # 置信度小标签
    p = doc.add_paragraph()
    p.add_run("结论置信度: ").bold = True
    p.add_run(concl["confidence"])

    # 方法学说明 — 收成一行脚注式,不展开
    doc.add_paragraph()  # 空行
    p = doc.add_paragraph()
    run = p.add_run(
        "—— 方法学说明: 本次基于简化 SBAS 反演,未接入大气校正,"
        "绝对速率不确定度约 ±5 mm/yr;若需毫米级精度,详见 GACOS 接入方案。"
    )
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    doc.save(str(output_path))
    return output_path


def _resolve_files_url(url: str | None) -> Path | None:
    if not url or not url.startswith("/files/"):
        return None
    base = Path("/opt/deepexplor-services/geo-insar/downloads").resolve()
    rel = url[len("/files/"):]
    return (base / rel).resolve()
